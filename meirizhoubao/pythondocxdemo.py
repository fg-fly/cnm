#coding=utf-8
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt


# reload(sys)  # reload 才能调用 setdefaultencoding 方法
# sys.setdefaultencoding('utf-8')
document = Document('d.docx')
document.styles['Normal'].font.name = u'微软雅黑'
document.styles['Normal'].font.size = Pt(14)
document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
styles = document.styles

#选取 style，并设置 style 中的段落格式
stylep = styles['Normal']
para_format = stylep.paragraph_format
para_format.left_indent = Pt(0)
para_format.widow_control = True

styler = styles['Normal']
para_format = styler.paragraph_format
para_format.left_indent = Pt(20)
para_format.widow_control = True


p1 = document.add_paragraph(u'一、系统服务器设备使用率检查',style=stylep)
#公文系统服务器检查
p2 = document.add_paragraph(u'（1）公文系统服务器检查',style=styler)
p2.space_before = Pt(15)
p3 = document.add_paragraph(u'公文系统:51.6.240.1')
p3.space_before = Pt(15)
p4 = document.add_paragraph(u'检查方式:通过堡垒机远程登录到服务器')
p5 = document.add_paragraph(u'检查频次:每日三次')
p6 = document.add_paragraph(u'检查内容:cpu使用率、内存使用率、磁盘使用率')
p7 = document.add_paragraph(u'检查结果如下:')
p8 = document.add_paragraph(u'检查方式:通过堡垒机远程登录到服务器')
table1 = document.add_table(6,4,style='Table Grid')
#header
cells = table1.rows[0].cells
cells[0].text=u"日期"
cells[1].text=u"CPU使用率"
cells[2].text=u"内存使用率"
cells[3].text=u"磁盘使用率"

# table_run1 = table1.cell(0,0).paragraphs[0].add_run('test')
p9 = document.add_paragraph(u'结果说明:CPU使用率≦50%属于正常范围，内存使用率≦80%属于正常范围，磁盘使用≦80%属于正常范围。')
#安邮系统服务器检查
p10 = document.add_paragraph(u'（1）公文系统服务器检查')
p2.space_before = Pt(15)
p11 = document.add_paragraph(u'公文系统:51.6.240.1')
p3.space_before = Pt(15)
p12 = document.add_paragraph(u'检查方式:通过堡垒机远程登录到服务器')
p13 = document.add_paragraph(u'检查频次:每日三次')
p14 = document.add_paragraph(u'检查内容:cpu使用率、内存使用率、磁盘使用率')
p15 = document.add_paragraph(u'检查结果如下:')
p16 = document.add_paragraph(u'检查方式:通过堡垒机远程登录到服务器')
table2 = document.add_table(6,4,style='Table Grid')
cells = table2.rows[0].cells
cells[0].text=u"日期"
cells[1].text=u"CPU使用率"
cells[2].text=u"内存使用率"
cells[3].text=u"磁盘使用率"

# table_run2 = table2.cell(0,0).paragraphs[0].add_run('test')
p17 = document.add_paragraph(u'结果说明:CPU使用率≦50%属于正常范围，内存使用率≦80%属于正常范围，磁盘使用≦80%属于正常范围。')
#门户系统服务器检查
p18 = document.add_paragraph(u'（1）公文系统服务器检查')
p2.space_before = Pt(15)
p18 = document.add_paragraph(u'公文系统:51.6.240.1')
p3.space_before = Pt(15)
p19 = document.add_paragraph(u'检查方式:通过堡垒机远程登录到服务器')
p20 = document.add_paragraph(u'检查频次:每日三次')
p21 = document.add_paragraph(u'检查内容:cpu使用率、内存使用率、磁盘使用率')
p22 = document.add_paragraph(u'检查结果如下:')
p23 = document.add_paragraph(u'检查方式:通过堡垒机远程登录到服务器')
table3 = document.add_table(6,4,style='Table Grid')
cells = table3.rows[0].cells
cells[0].text=u"日期"
cells[1].text=u"CPU使用率"
cells[2].text=u"内存使用率"
cells[3].text=u"磁盘使用率"

# table_run3 = table3.cell(0,0).paragraphs[0].add_run('test')
p24 = document.add_paragraph(u'结果说明:CPU使用率≦50%属于正常范围，内存使用率≦80%属于正常范围，磁盘使用≦80%属于正常范围。')

#每日更新内容
p25 = document.add_paragraph(u'二、门户网站更新工作')
p0 = document.add_paragraph(u'10月21日至10月25日更新汇总表:')
#更新汇总表
table_total = document.add_table(6,7,style='Table Grid')
cells = table_total.rows[0].cells
cells[0].text=u"星期"
cells[1].text=u"领导活动栏目"
cells[2].text=u"政务资源栏目"
cells[3].text=u"网络聚焦栏目"
cells[4].text=u"媒体报刊栏目"
cells[5].text=u"谈治国理政栏目"
cells[6].text=u'视频点播'

p26 = document.add_paragraph(u'1、领导活动栏目')
#领导活动表格
table_ldhd = document.add_table(6,2,style='Table Grid')
cells = table_total.rows[0].cells
# cells[0].text="星期"
cells[1].text=u"星期一"
cells[2].text=u"星期二"
cells[3].text=u"星期三"
cells[4].text=u"星期四"
cells[5].text=u"星期五"
table_ldhd.cell(2,0).text = u'领导活动'
p27 = document.add_paragraph(u'2、政务资源栏目')
#政务资源表格
table_zwzy = document.add_table(6,6,style='Table Grid')
cells = table_zwzy.rows[0].cells
# cells[0].text="星期"
cells[1].text=u"星期一"
cells[2].text=u"星期二"
cells[3].text=u"星期三"
cells[4].text=u"星期四"
cells[5].text=u"星期五"
table_ldhd.cell(2,0).text = u'区政府发文'
p28 = document.add_paragraph(u'3、网络聚焦栏目')
#网络聚焦表格
table_wwjj = document.add_table(6,6,style='Table Grid')
cells = table_wwjj.rows[0].cells
cells[0].text=u"栏目"
cells[1].text=u"星期一"
cells[2].text=u"星期二"
cells[3].text=u"星期三"
cells[4].text=u"星期四"
cells[5].text=u"星期五"
table_wwjj.cell(1,1).text = u'国内新闻'
table_wwjj.cell(2,1).text = u'国际新闻'
table_wwjj.cell(3,1).text = u'财经新闻'
table_wwjj.cell(4,1).text = u'健康新闻'
table_wwjj.cell(5,1).text = u'北京新闻'
p29 = document.add_paragraph(u'4、媒体报刊栏目')
#媒体报刊表格
table_mtbk = document.add_table(6,6,style='Table Grid')
cells = table_mtbk.rows[0].cells
cells[0].text=u"栏目"
cells[1].text=u"星期一"
cells[2].text=u"星期二"
cells[3].text=u"星期三"
cells[4].text=u"星期四"
cells[5].text=u"星期五"
table_mtbk.cell(1,1).text = u'人民日报'
table_mtbk.cell(2,1).text = u'北京日报'
table_mtbk.cell(3,1).text = u'海淀报'

#视频点播
ps = document.add_paragraph(u'5、视频点播')
table_spdb = document.add_table(6,6,style='Table Grid')
cells = table_spdb.rows[0].cells
cells[0].text=u"栏目"
cells[1].text=u"星期一"
cells[2].text=u"星期二"
cells[3].text=u"星期三"
cells[4].text=u"星期四"
cells[5].text=u"星期五"
table_mtbk.cell(1,1).text = u'新闻联播'
table_mtbk.cell(2,1).text = u'北京新闻'
table_mtbk.cell(3,1).text = u'海淀新闻'

p30 = document.add_paragraph(u'5、习近平谈治国理政栏目更新x条')

#电话接待工作
document.add_paragraph(u'三、系统保障与运维工作')
p31 = document.add_paragraph(u'1、电话解决问题2次')
p32 = document.add_paragraph(u'电话记录表:')
table_phone = document.add_table(3,6,style='Table Grid')

#下周工作安排
p33 = document.add_paragraph()


#设置缩进
# p1.space_before = Pt(10)
document.save('destination.docx')

