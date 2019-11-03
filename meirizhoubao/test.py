# from docx.enum.style import WD_STYLE_TYPE
# from docx import *
#
# document = Document()
# styles = document.styles
#
# for s in styles:
#     if s.type == WD_STYLE_TYPE.PARAGRAPH:
#         document.add_paragraph('Paragraph style is : '+ s.name, style = s)
#
# document.save('para_style.docx')
# -*- coding: UTF-8 -*-
import sys

from copy import deepcopy

import xlwings
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def get_excel_date(filename):
    '''
    获得excel里的所有内容，返回list
    :param filename:  excel路径
    :return: list[list[]]
    '''
    app = xlwings.App(visible=False, add_book=True)
    app.display_alerts = False
    app.screen_updating = False
    wb = app.books.open(filename)
    sht = wb.sheets[0]
    rng = sht.range('A1')
    # 把excel里的数据读取成 年-月-日 时:分:秒的格式
    my_date_handler = lambda year, month, day, hour, minute, second, **kwargs: "%04i-%02i-%02i %02i:%02i:%02i" % (
        year, month, day, hour, minute, second)
    # 取出所有内容,这里用ig这个变量，是为了庆祝I.G获得LOL S8赛季总冠军
    ig = rng.current_region.options(index=False, numbers=int, empty='N/A', dates=my_date_handler)
    result = ig.value
    wb.close()
    app.quit()
    return result

def delete_table_with_title(document,expect_text):
    allTables = document.tables
    for activeTable in allTables:
        if activeTable.cell(0, 0).paragraphs[0].text == expect_text:
            print('删除成功')
            activeTable._element.getparent().remove(activeTable._element)

def insert_table_after_text(file_name,excel_name,expect_text):
    document = Document(file_name)
    # 因为docx读出来的都是unicode类型的，所以我们要用unicode类型的进行查找
    expect_text=expect_text.decode('utf-8')
    delete_table_with_title(document,expect_text)
    target = None
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    if target is not None:
        records = get_excel_date(excel_name)
        # 获得excel数据的栏数，初始化一个空的table
        col = len(records[0])
        table = document.add_table(rows=1, cols=col)
        table.style = 'Table Grid'
        # 给table加一个表头，并且合并第一栏
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9E2F3"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        table.rows[0].cells[0].text=expect_text
        table_row=table.rows[0]
        first=table_row.cells[0]
        end=table_row.cells[-1]
        first.merge(end)
        # 合并结束，开始把excel里的内容添加到table里
        for tr_list in records:
            row_cells = table.add_row().cells
            index = 0
            for td_list in tr_list:
                row_cells[index].text = td_list
                index = index + 1
        # 把添加的table移动到指定的位置
        move_table_after(table, target)
        # 保存
    document.save(file_name)
if __name__ == '__main__':
    insert_table_after_text('demo2.docx', 'demo.xlsx',"BUG情况表")